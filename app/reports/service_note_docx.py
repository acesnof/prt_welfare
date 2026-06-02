import os
from datetime import date

from docx import Document

MONTH_ABBR_EN = ["JAN", "FEB", "MAR", "APR", "MAY", "JUN", "JUL", "AUG", "SEP", "OCT", "NOV", "DEC"]
MONTH_FULL_EN = [
    "JANUARY", "FEBRUARY", "MARCH", "APRIL", "MAY", "JUNE",
    "JULY", "AUGUST", "SEPTEMBER", "OCTOBER", "NOVEMBER", "DECEMBER"
]


def formatar_data_militar(data_obj):
    return f"{data_obj.day:02d}{MONTH_ABBR_EN[data_obj.month - 1]}{str(data_obj.year)[-2:]}"


def mes_ano_en(ano, mes):
    return f"{MONTH_FULL_EN[mes - 1]} {ano}"


def _procurar_template(docs_dir):
    candidatos = [
        os.path.join(docs_dir, "reimbursment_service_note.docx"),
        os.path.join(docs_dir, "reimbursement_service_note.docx"),
        os.path.join(docs_dir, "reimbursment_service_note"),
        os.path.join(docs_dir, "reimbursement_service_note"),
    ]
    for caminho in candidatos:
        if os.path.exists(caminho):
            return caminho
    return None


def _replace_in_paragraph(paragraph, replacements):
    texto_original = paragraph.text
    if not texto_original:
        return

    texto_novo = texto_original
    alterado = False
    for chave, valor in replacements.items():
        if chave in texto_novo:
            texto_novo = texto_novo.replace(chave, valor or "")
            alterado = True

    if not alterado:
        return

    # Mantém o estilo do parágrafo e o estilo do primeiro run.
    if paragraph.runs:
        primeiro_run = paragraph.runs[0]
        for run in paragraph.runs:
            run.text = ""
    else:
        primeiro_run = paragraph.add_run()

    partes = texto_novo.split("\n")
    for idx, parte in enumerate(partes):
        if idx > 0:
            primeiro_run.add_break()
        primeiro_run.add_text(parte)


def _replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            for nested in cell.tables:
                _replace_in_table(nested, replacements)


def substituir_placeholders_docx(template_path, destino, replacements):
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        _replace_in_table(table, replacements)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in section.header.tables:
            _replace_in_table(table, replacements)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in section.footer.tables:
            _replace_in_table(table, replacements)

    doc.save(destino)


def gerar_service_note(docs_dir, destino, ano, mes, dates_cohesion, individual_cohesion, chief_of_staff_name):
    template = _procurar_template(docs_dir)
    if not template:
        raise FileNotFoundError("docs/reimbursment_service_note.docx")

    hoje = date.today()
    replacements = {
        "DATA_HOJE": formatar_data_militar(hoje),
        "ANO_ATUAL": str(ano),
        "MES_ANO": mes_ano_en(ano, mes),
        "DATES_COESION": dates_cohesion or "",
        "INDIVIDUAL_COESION": individual_cohesion or "",
        "CHIEF_OF_STAF_NAME": chief_of_staff_name or "",
        # Também suporta a grafia correta, caso o template a use.
        "CHIEF_OF_STAFF_NAME": chief_of_staff_name or "",
    }

    substituir_placeholders_docx(template, destino, replacements)
